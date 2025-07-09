# Sebastian Heredia
# DyHeredia@mednet.ucla.edu
# July 3, 2025


import streamlit as st                      # Loads Streamlit GUI 
from docx import Document                   # For working with .docx document
from docx.shared import Pt, RGBColor        # For setting font size + Color
from io import BytesIO                      # To hold input/output memory for download
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn                 # To enable .font.name update
from docx.oxml import OxmlElement           # Used to manually adjust width

import firebase_admin                       # Sync with Firebase
from firebase_admin import credentials, db


# CCS: Widen sidebar to fit text on laptop, keep collapsable for mobile
st.markdown(
    """
    <style>
        @media (min-width: 768px) {
            /* On tablets and desktops, widen the sidebar */
            [data-testid="stSidebar"][aria-expanded="true"] {
                width: 400px !important;
                min-width: 400px !important;
                max-width: 400px !important;
            }

            [data-testid="stSidebar"][aria-expanded="true"] > div:first-child {
                width: 100% !important;
            }

            [data-testid="stSidebar"][aria-expanded="false"] {
                width: 0 !important;
                min-width: 0 !important;
                max-width: 0 !important;
            }

            section[data-testid="stSidebar"] {
                flex-shrink: 0 !important;
            }
        }
    </style>
    """,
    unsafe_allow_html=True
)


# Initialize Firebase once
if not firebase_admin._apps:
   cred = credentials.Certificate({
       "type": st.secrets["firebase"]["type"],
       "project_id": st.secrets["firebase"]["project_id"],
       "private_key_id": st.secrets["firebase"]["private_key_id"],
       "private_key": st.secrets["firebase"]["private_key"].replace('\\n', '\n'),
       "client_email": st.secrets["firebase"]["client_email"],
       "client_id": st.secrets["firebase"]["client_id"],
       "auth_uri": st.secrets["firebase"]["auth_uri"],
       "token_uri": st.secrets["firebase"]["token_uri"],
       "auth_provider_x509_cert_url": st.secrets["firebase"]["auth_provider_x509_cert_url"],
       "client_x509_cert_url": st.secrets["firebase"]["client_x509_cert_url"],
       "universe_domain": st.secrets["firebase"]["universe_domain"]
   })
   firebase_admin.initialize_app(cred, {
       'databaseURL': 'https://dgsom-cv-formatter-default-rtdb.firebaseio.com/'
   })


def save_to_firebase():
   db.reference("cvs").set(st.session_state.all_cvs)


if "all_cvs" not in st.session_state:
   firebase_data = db.reference("cvs").get()
   if firebase_data:
       st.session_state.all_cvs = firebase_data
   else:
       st.session_state.all_cvs = {}


if "current_cv" not in st.session_state:
   st.session_state.current_cv = "Default_CV"


if st.session_state.current_cv not in st.session_state.all_cvs:
   st.session_state.all_cvs[st.session_state.current_cv] = {
       "BUSINESS INFORMATION": [
           {"name": "", "position": "", "order": 1, "company_name": "",
            "last_updated": "", "business_address": "", "business_phone": "",
            "email": ""}
       ],
       "EDUCATION": [
           {"degree": "", "year": "", "order": 1, "school": ""}
       ]
   }


st.session_state.cv_data = st.session_state.all_cvs[st.session_state.current_cv]


st.title("UCLA DGSOM CV Organizer")


# Sidebar Login placeholder
st.sidebar.header("Login")
st.sidebar.button("[Insert Login Here]")
st.sidebar.markdown("---")


# CV Manager Sidebar
st.sidebar.header("CV Manager")
cv_list = list(st.session_state.all_cvs.keys())
selected_cv = st.sidebar.selectbox("Select CV", cv_list, index=cv_list.index(st.session_state.current_cv))


if selected_cv != st.session_state.current_cv:
   st.session_state.current_cv = selected_cv
   st.session_state.cv_data = st.session_state.all_cvs[selected_cv]
   st.rerun()


new_cv_name = st.sidebar.text_input("New CV name (no spaces):")
if st.sidebar.button("Create New CV"):
   if new_cv_name and new_cv_name not in st.session_state.all_cvs:
       st.session_state.all_cvs[new_cv_name] = {
           "BUSINESS INFORMATION": [
               {"name": "", "position": "", "order": 1, "company_name": "",
                "last_updated": "", "business_address": "", "business_phone": "",
                "email": ""}
           ],
           "EDUCATION": [
               {"degree": "", "year": "", "order": 1, "school": ""}
           ]
       }
       st.session_state.current_cv = new_cv_name
       st.session_state.cv_data = st.session_state.all_cvs[new_cv_name]
       save_to_firebase()
       st.success(f"New CV '{new_cv_name}' created.")
       st.rerun()
   else:
       st.sidebar.warning("Enter a unique CV name.")


# Rename CV
st.sidebar.header("Rename CV")
with st.sidebar.expander("Rename CV"):
   new_CVname = st.text_input("Enter new name", value=st.session_state.current_cv)
   if st.button("Rename"):
       if new_CVname and new_CVname != st.session_state.current_cv:
           cvs = st.session_state.all_cvs
           current = st.session_state.current_cv
           if new_CVname in cvs:
               st.warning("A CV with that name already exists.")
           else:
               cvs[new_CVname] = cvs.pop(current)
               st.session_state.current_cv = new_CVname
               save_to_firebase()
               st.success(f"Renamed to '{new_CVname}'")
               st.rerun()


# Delete CV with confirmation
delete_cv_flag_key = f"delete_confirm_cv_{st.session_state.current_cv}"
st.sidebar.markdown("---")
st.sidebar.header("Delete CV")


if delete_cv_flag_key not in st.session_state:
   st.session_state[delete_cv_flag_key] = False


if not st.session_state[delete_cv_flag_key]:
   if st.sidebar.button(f'Delete "{st.session_state.current_cv}"'):
       st.session_state[delete_cv_flag_key] = True
       st.rerun()
else:
   st.sidebar.write(f"Are you sure you want to delete '{st.session_state.current_cv}'?")
   confirm = st.sidebar.button("Yes, delete")
   cancel = st.sidebar.button("Cancel")
   if confirm:
       st.session_state.all_cvs.pop(st.session_state.current_cv, None)
       save_to_firebase()
       if "Default_CV" in st.session_state.all_cvs:
           st.session_state.current_cv = "Default_CV"
       elif len(st.session_state.all_cvs) > 0:
           st.session_state.current_cv = list(st.session_state.all_cvs.keys())[0]
       else:
           st.session_state.all_cvs["Default_CV"] = {
               "BUSINESS INFORMATION": [
                   {"name": "", "position": "", "order": 1, "company_name": "",
                    "last_updated": "", "business_address": "", "business_phone": "",
                    "email": ""}
               ],
               "EDUCATION": [
                   {"degree": "", "year": "", "order": 1, "school": ""}
               ]
           }
           st.session_state.current_cv = "Default_CV"
       st.session_state.cv_data = st.session_state.all_cvs[st.session_state.current_cv]
       st.session_state.pop(delete_cv_flag_key, None)
       st.rerun()
   if cancel:
       st.session_state[delete_cv_flag_key] = False
       st.rerun()


# BUSINESS INFORMATION Section
st.header("BUSINESS INFORMATION")


for i, entry in enumerate(st.session_state.cv_data.get("BUSINESS INFORMATION", [])):
   name_display = entry["name"].strip() or ""
   with st.expander(f"Entry {i+1}: {name_display}"):
       with st.form(f"biz_form_{st.session_state.current_cv}_{i}"):
           new_name = st.text_input("Name", value=entry.get("name", ""), key=f"name_{st.session_state.current_cv}_{i}")
           new_position = st.text_input("Degree(s)", value=entry.get("position", ""), key=f"position_{st.session_state.current_cv}_{i}")
           new_last_updated = st.text_input("Last Updated", value=entry.get("last_updated", ""), key=f"last_updated_{st.session_state.current_cv}_{i}")
           new_company = st.text_input("Company Name", value=entry.get("company_name", ""), key=f"company_{st.session_state.current_cv}_{i}")
           new_address = st.text_input("Business Address", value=entry.get("business_address", ""), key=f"address_{st.session_state.current_cv}_{i}")
           new_phone = st.text_input("Business Phone", value=entry.get("business_phone", ""), key=f"phone_{st.session_state.current_cv}_{i}")
           new_email = st.text_input("Email", value=entry.get("email", ""), key=f"email_{st.session_state.current_cv}_{i}")
           submitted = st.form_submit_button("Update Entry")
           if submitted:
               entry["name"] = new_name
               entry["position"] = new_position
               entry["last_updated"] = new_last_updated
               entry["company_name"] = new_company
               entry["business_address"] = new_address
               entry["business_phone"] = new_phone
               entry["email"] = new_email
               save_to_firebase()
               st.rerun()


# EDUCATION Section
st.header("EDUCATION")


for i, entry in enumerate(st.session_state.cv_data["EDUCATION"]):
   degree_display = entry["degree"].strip() or ""
   year_display = entry["year"].strip() or ""
   with st.expander(f"Entry {i+1}: {year_display} {degree_display}"):
       with st.form(f"form_{st.session_state.current_cv}_{i}"):
           new_degree = st.text_input("Degree", value=entry.get("degree", ""), key=f"degree_{st.session_state.current_cv}_degree_{i}")
           new_year = st.text_input("Year (Ex: 2000-2004)", value=entry.get("year", ""), key=f"year_{st.session_state.current_cv}_year_{i}")
           new_school = st.text_input("University Name", value=entry.get("school", ""), key=f"school_{st.session_state.current_cv}_school_{i}")
           submitted = st.form_submit_button("Update Entry")


           flag_key = f"delete_confirm_{st.session_state.current_cv}_{i}"
           if flag_key not in st.session_state:
               st.session_state[flag_key] = False


           if not st.session_state[flag_key]:
               delete_clicked = st.form_submit_button("Delete Entry")
               if delete_clicked:
                   st.session_state[flag_key] = True
                   st.rerun()
           else:
               st.write("Are you sure you want to delete this entry?")
               confirm = st.form_submit_button("Yes, delete")
               cancel = st.form_submit_button("Cancel")
               if confirm:
                   st.session_state.cv_data["EDUCATION"].pop(i)
                   st.session_state.pop(flag_key, None)
                   save_to_firebase()
                   st.rerun()
               if cancel:
                   st.session_state[flag_key] = False
                   st.rerun()


           if submitted:
               entry["degree"] = new_degree
               entry["year"] = new_year
               entry["school"] = new_school
               save_to_firebase()
               st.rerun()


if st.button("Add Entry", key="add_entry_EDUCATION"):
   st.session_state.cv_data.setdefault("EDUCATION", []).append({
       "degree": "",
       "year": "",
       "school": "",
       "order": len(st.session_state.cv_data.get("EDUCATION", [])) + 1
   })
   save_to_firebase()
   st.rerun()


st.header("LECTURES AND PRESENTATIONS")
st.header("BIBLIOGRAPHY")




# Generate .docx and download button
def generate_docx(data):
    doc = Document()

    # Pull first business info entry for name/degree
    biz_info = data.get("BUSINESS INFORMATION", [])
    name = biz_info[0]["name"].strip() if biz_info else ""
    degree = biz_info[0].get("position", "").strip() if biz_info else ""
    updated = biz_info[0].get("last_updated", "").strip() if biz_info else ""

    name_line = f"{name}, {degree}".strip().rstrip(",")
    last_updated_line = f"{updated}".strip().rstrip(",")

    lines = [
        {"text": name_line, "bold": True, "italic": False},
        {"text": "Curriculum Vitae", "bold": True, "italic": False},
        {"text": last_updated_line, "bold": False, "italic": True},
    ]

    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line["text"])
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        run.font.bold = line["bold"]
        run.font.italic = line["italic"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    # BUSINESS INFORMATION Header
    biz_para = doc.add_paragraph()
    biz_run = biz_para.add_run("BUSINESS INFORMATION")
    biz_run.font.name = 'Times New Roman'
    biz_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    biz_run.font.size = Pt(12)
    biz_run.font.bold = True
    biz_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    biz_para.paragraph_format.space_after = Pt(3)
    biz_para.paragraph_format.space_before = Pt(16)

    # BUSINESS INFORMATION entries
    for entry in sorted(data.get("BUSINESS INFORMATION", []), key=lambda x: x["order"]):
        updated = entry.get("last_updated", "").strip() or "Missing Updated Date"
        company = entry.get("company_name", "").strip() or "Missing Company"
        address = entry.get("business_address", "").strip() or "Missing Address"
        phone = entry.get("business_phone", "").strip() or "Missing Phone"
        email = entry.get("email", "").strip() or "Missing Email"

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        lines = [updated, company, address, f"Phone: {phone}", f"Email: {email}"]

        for i, line in enumerate(lines):
            if "Missing" in line:
                before, after = line.split("Missing", 1)
                para.add_run(before).font.name = 'Times New Roman'
                red_run = para.add_run("Missing" + after)
                red_run.font.name = 'Times New Roman'
                red_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                para.add_run(line).font.name = 'Times New Roman'
            if i < len(lines) - 1:
                para.add_run().add_break()

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    # EDUCATION header
    edu_para = doc.add_paragraph()
    edu_run = edu_para.add_run("EDUCATION")
    edu_run.font.name = 'Times New Roman'
    edu_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    edu_run.font.size = Pt(12)
    edu_run.font.bold = True
    edu_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    edu_para.paragraph_format.space_after = Pt(3)
    edu_para.paragraph_format.space_before = Pt(16)

    # EDUCATION table with 3 columns
    table = doc.add_table(rows=0, cols=3)
    col_widths = [1.15, 1.50, 3.85]  # inches

    for entry in sorted(data["EDUCATION"], key=lambda x: x["order"]):
        year = entry.get("year", "").strip() or "Year Missing"
        degree = entry.get("degree", "").strip() or "Degree Missing"
        school = entry.get("school", "").strip() or "School Missing"

        if not any([year, degree, school]):
            continue

        row_cells = table.add_row().cells
        row_data = [year, degree, school]

        for i, (cell, text) in enumerate(zip(row_cells, row_data)):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            if "Missing" in text:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

            # Set column width (in twips)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[i] * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # Left indent for all table paragraphs
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = Pt(16)

    # Page break before bibliography
    page_break_para = doc.add_paragraph()
    page_break_run = page_break_para.add_run()
    page_break_run.add_break(WD_BREAK.PAGE)

    # BIBLIOGRAPHY header
    bib_para = doc.add_paragraph()
    bib_run = bib_para.add_run("BIBLIOGRAPHY")
    bib_run.font.name = 'Times New Roman'
    bib_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    bib_run.font.size = Pt(12)
    bib_run.font.bold = True
    bib_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bib_para.paragraph_format.space_after = Pt(3)
    bib_para.paragraph_format.space_before = Pt(16)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ⬇️ Streamlit download button
st.download_button(
    label="Download CV",
    data=generate_docx(st.session_state.cv_data),
    file_name=f"{st.session_state.current_cv}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
